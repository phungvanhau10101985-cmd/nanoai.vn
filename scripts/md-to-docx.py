#!/usr/bin/env python3
"""Convert project markdown ops doc to .docx (python-docx)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def add_rich_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        else:
            p.add_run(part)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [strip_md_inline(c.strip()) for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|\s*$", line.strip()))


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_paragraph(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for c, h in enumerate(headers):
                cell = table.rows[0].cells[c]
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            for r_idx, row in enumerate(rows):
                for c_idx, val in enumerate(row):
                    if c_idx < len(table.rows[r_idx + 1].cells):
                        table.rows[r_idx + 1].cells[c_idx].text = val
            doc.add_paragraph()
            continue

        if stripped.startswith("### "):
            doc.add_heading(strip_md_inline(stripped[4:]), level=3)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(strip_md_inline(stripped[3:]), level=2)
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(strip_md_inline(stripped[2:]), level=1)
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m:
            add_rich_paragraph(doc, f"{m.group(1)}. {m.group(2)}", style="List Number")
            i += 1
            continue

        if stripped.startswith("- "):
            add_rich_paragraph(doc, stripped[2:], style="List Bullet")
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        add_rich_paragraph(doc, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else root / "docs" / "khoi-phuc-web-vps.md"
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix(".docx")
    convert(md, out)


if __name__ == "__main__":
    main()
